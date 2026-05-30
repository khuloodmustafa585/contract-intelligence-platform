from concurrent.futures import ThreadPoolExecutor, TimeoutError as FutureTimeout
import io
import os

import pytesseract

from app.core.logging import app_logger
from app.utils.text_cleaner import clean_text

PARSER_TIMEOUT_SECONDS = 45
MIN_EXTRACTED_TEXT_CHARS = 10
DEFAULT_TESSERACT_CMD = r"C:\Program Files\Tesseract-OCR\tesseract.exe"


if os.path.exists(DEFAULT_TESSERACT_CMD):
    pytesseract.pytesseract.tesseract_cmd = DEFAULT_TESSERACT_CMD


def _ensure_non_empty_text(text: str, source: str) -> str:
    text = (text or "").strip()
    if len(text) < MIN_EXTRACTED_TEXT_CHARS:
        raise ValueError(f"{source} produced no usable text")
    return text


def _ocr_image(image, source: str) -> str:
    try:
        text = pytesseract.image_to_string(image)
        app_logger.info("%s OCR extracted %s chars", source, len(text or ""))
        return text.strip()
    except pytesseract.TesseractNotFoundError as exc:
        raise RuntimeError("Tesseract OCR executable was not found") from exc
    except Exception as exc:
        raise RuntimeError(f"{source} OCR failed: {exc}") from exc


def _ocr_pdf_pages(document) -> str:
    from PIL import Image

    page_texts: list[str] = []
    zoom = 2

    for page_index, page in enumerate(document, start=1):
        import fitz

        pixmap = page.get_pixmap(matrix=fitz.Matrix(zoom, zoom), alpha=False)
        image = Image.open(io.BytesIO(pixmap.tobytes("png")))
        try:
            page_text = _ocr_image(image, f"PDF page {page_index}")
        finally:
            image.close()
        if page_text:
            page_texts.append(page_text)

    return "\n".join(page_texts).strip()


def extract_text_from_pdf(file_path: str) -> tuple[str, str, bool]:
    try:
        import fitz

        with fitz.open(file_path) as document:
            page_texts = [page.get_text("text") for page in document]
            text = "\n".join(page_texts).strip()
            app_logger.info(
                "PDF direct extraction: file=%s pages=%s chars=%s",
                os.path.basename(file_path),
                document.page_count,
                len(text),
            )

            if len(text) >= MIN_EXTRACTED_TEXT_CHARS:
                return text, "pdf", False

            app_logger.warning(
                "PDF direct extraction empty; falling back to OCR: file=%s pages=%s",
                os.path.basename(file_path),
                document.page_count,
            )
            ocr_text = _ocr_pdf_pages(document)
            return _ensure_non_empty_text(ocr_text, "PDF OCR"), "pdf_ocr", True
    except Exception as e:
        raise Exception(f"PDF parsing failed: {str(e)}")


def extract_text_from_docx(file_path: str) -> tuple[str, str, bool]:
    try:
        import docx

        document = docx.Document(file_path)
        parts: list[str] = []

        for para in document.paragraphs:
            if para.text:
                parts.append(para.text)

        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))

        text = "\n".join(parts).strip()
        app_logger.info("DOCX extraction: file=%s chars=%s", os.path.basename(file_path), len(text))
        return _ensure_non_empty_text(text, "DOCX extraction"), "docx", False
    except Exception as e:
        raise Exception(f"DOCX parsing failed: {str(e)}")


def extract_text_from_image(file_path: str) -> tuple[str, str, bool]:
    try:
        from PIL import Image

        image = Image.open(file_path)
        try:
            text = _ocr_image(image, f"Image {os.path.basename(file_path)}")
        finally:
            image.close()
        return _ensure_non_empty_text(text, "Image OCR"), "ocr", True
    except Exception as e:
        raise Exception(f"Image parsing failed: {str(e)}")


def _extract_text(file_path: str, file_type: str) -> tuple[str, str, bool]:
    if file_type == "pdf":
        return extract_text_from_pdf(file_path)
    elif file_type == "docx":
        return extract_text_from_docx(file_path)
    elif file_type in ["jpg", "jpeg", "png"]:
        return extract_text_from_image(file_path)
    else:
        raise ValueError("Unsupported file type for text extraction")


def extract_text(file_path: str, file_type: str, timeout_seconds: int = PARSER_TIMEOUT_SECONDS) -> dict:
    executor = ThreadPoolExecutor(max_workers=1)
    try:
        future = executor.submit(_extract_text, file_path, file_type)
        try:
            raw_text, parse_method, ocr_used = future.result(timeout=timeout_seconds)
        except FutureTimeout:
            future.cancel()
            app_logger.exception("Text extraction timed out: file=%s type=%s", file_path, file_type)
            raise TimeoutError("Text extraction timed out")
    finally:
        executor.shutdown(wait=False, cancel_futures=True)

    cleaned_text = clean_text(raw_text or "")
    app_logger.info(
        "Text extraction result: file=%s type=%s method=%s raw_len=%s cleaned_len=%s ocr_used=%s",
        os.path.basename(file_path),
        file_type,
        parse_method,
        len(raw_text or ""),
        len(cleaned_text or ""),
        ocr_used,
    )
    _ensure_non_empty_text(raw_text, "Raw text extraction")
    _ensure_non_empty_text(cleaned_text, "Cleaned text generation")

    return {
        "raw_text": raw_text,
        "cleaned_text": cleaned_text,
        "ocr_used": ocr_used,
        "parse_method": parse_method,
    }
